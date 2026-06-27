"""
Multi-Agent Executor — Uses Claude with MCP tools to execute approved resolution steps.

Flow:
1. Receives approved resolution steps for a case
2. Claude acts as an autonomous agent, calling MCP tools to execute each step
3. Logs every action taken on the case
4. Returns execution results
"""

import json
import re

import anthropic
import structlog
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor

from config import settings
from salesforce_client import SalesforceClient
from vector_store import VectorStore
from mcp_tools import MCPToolExecutor, TOOL_DEFINITIONS

log = structlog.get_logger()

EXECUTOR_SYSTEM_PROMPT = """\
You are an AI Case Resolution Agent executing approved resolution steps on Salesforce.

You have been given a set of resolution steps that a human agent has approved.
Execute each step in order using the tools available to you.

RULES:
- Execute EVERY approved step. Do not skip any.
- Use the appropriate tool for each action.
- Log each action by adding a case comment before executing it.
- If a step fails, add a comment explaining the failure and continue with the next step.
- After all steps are complete, close the case if the steps indicate it should be closed.
- Be precise with email content — use the exact templates provided in the resolution.
- Always add a final summary comment listing all actions taken.
"""

REVISE_PROMPT = """\
You are an AI-powered Salesforce Case Resolution Engine.

A human agent has reviewed your previous resolution proposal and provided feedback.
Revise your resolution based on their feedback.

=== CASE DETAILS ===
Case ID: {case_id}
Case Number: {case_number}
Subject: {subject}
Description: {description}
Type: {case_type}

=== YOUR PREVIOUS RESOLUTION ===
{previous_resolution}

=== AGENT FEEDBACK ===
Action: {agent_action}
Comments: {agent_comments}

=== HISTORICAL KNOWLEDGE ===
{knowledge_articles}

=== COMPANY GUIDELINES ===
{company_knowledge}

Based on the agent's feedback, provide a REVISED resolution.
If the agent said "Partially Correct", keep the correct parts and fix the incorrect parts.
If the agent said "Not Right", rethink the resolution entirely.

IMPORTANT RULES:
- Be CONCISE. No filler text.
- resolution_steps: Each step under 15 words. Short actionable instructions only.
- resolution_notes: 2-3 sentences max.
- customer_email_body: Under 100 words.
- All other fields: 1 sentence max.

Respond in the following JSON format ONLY (no markdown, no extra text):
{{
    "issue_summary": "...",
    "root_cause": "...",
    "recommended_action": "auto_resolve | escalate | route | request_info",
    "resolution_steps": ["short step 1", "short step 2"],
    "resolution_notes": "Brief resolution summary",
    "requires_human": true/false,
    "route_to_team": "team name or null",
    "customer_email_subject": "...",
    "customer_email_body": "...",
    "can_auto_close": true/false,
    "confidence_score": 0.0-1.0,
    "reasoning": "Brief reason"
}}
"""


class AgentExecutor:
    def __init__(self, sf_client: SalesforceClient, vector_store: VectorStore = None):
        self.sf = sf_client
        self.vector_store = vector_store
        self.tool_executor = MCPToolExecutor(sf_client)
        self.llm = anthropic.Anthropic(api_key=settings.anthropic_api_key)

    def execute_steps(self, case_id: str, resolution_steps: list[str], resolution_data: dict) -> dict:
        log.info("executing_resolution", case_id=case_id, steps=len(resolution_steps))

        messages = [
            {
                "role": "user",
                "content": (
                    f"Execute the following approved resolution steps for Case ID: {case_id}\n\n"
                    f"Case Details:\n"
                    f"  Subject: {resolution_data.get('issue_summary', '')}\n"
                    f"  Root Cause: {resolution_data.get('root_cause', '')}\n"
                    f"  Customer Email Subject: {resolution_data.get('customer_email_subject', '')}\n"
                    f"  Customer Email Body: {resolution_data.get('customer_email_body', '')}\n\n"
                    f"Steps to execute:\n"
                    + "\n".join(f"  {i+1}. {s}" for i, s in enumerate(resolution_steps))
                    + "\n\nExecute each step now using the available tools. "
                    + "Do NOT add case comments for individual steps — only use tools for actual actions."
                ),
            }
        ]

        execution_log = []
        failures = []
        max_iterations = 15

        try:
            for iteration in range(max_iterations):
                response = self.llm.messages.create(
                    model=settings.anthropic_model,
                    max_tokens=4096,
                    system=EXECUTOR_SYSTEM_PROMPT,
                    tools=TOOL_DEFINITIONS,
                    messages=messages,
                )

                if response.stop_reason == "end_turn":
                    final_text = ""
                    for block in response.content:
                        if block.type == "text":
                            final_text = block.text
                    execution_log.append({"type": "completion", "message": final_text})
                    break

                assistant_content = response.content
                messages.append({"role": "assistant", "content": assistant_content})

                tool_results = []
                for block in assistant_content:
                    if block.type == "tool_use":
                        log.info("executing_tool", tool=block.name, iteration=iteration)
                        result = self.tool_executor.execute_tool(block.name, block.input)
                        is_error = result.startswith("Error")
                        execution_log.append({
                            "type": "tool_call",
                            "tool": block.name,
                            "input": block.input,
                            "result": result,
                            "success": not is_error,
                        })
                        if is_error:
                            failures.append(f"{block.name}: {result}")
                        tool_results.append({
                            "type": "tool_result",
                            "tool_use_id": block.id,
                            "content": result,
                        })

                if tool_results:
                    messages.append({"role": "user", "content": tool_results})

        except Exception as e:
            failures.append(f"Execution error: {e}")
            log.error("execution_exception", case_id=case_id, error=str(e))

        # Build execution summary
        tool_calls = [e for e in execution_log if e["type"] == "tool_call"]
        successful = [e for e in tool_calls if e.get("success", True)]
        failed = [e for e in tool_calls if not e.get("success", True)]

        summary_lines = [f"Actions: {len(tool_calls)} total, {len(successful)} succeeded, {len(failed)} failed"]
        for e in tool_calls:
            status = "OK" if e.get("success", True) else "FAILED"
            summary_lines.append(f"  [{status}] {e['tool']}: {e['result'][:120]}")

        summary = "\n".join(summary_lines)

        # Comment: only final summary
        self.tool_executor.execute_tool("add_case_comment", {
            "case_id": case_id,
            "comment": f"[AI Agent] Execution Summary\n{summary}",
        })

        # Update AI fields based on success/failure
        if failures:
            failure_detail = "\n".join(failures)
            try:
                self.sf.update_case(case_id, {
                    "AI_Resolution_Status__c": "Execute Resolution",
                    "AI_Recommended_Action__c": f"EXECUTION FAILED\n{failure_detail}\n\nSummary:\n{summary}",
                })
            except Exception:
                pass

            log.warning("execution_partial_failure", case_id=case_id, failures=len(failures))
        else:
            # Build rich text notes with execution results
            steps_html = "".join(
                f'<li>{e["tool"]}: {e["result"][:150]}</li>' for e in tool_calls
            )
            rich_notes = (
                f"<h2>Execution Completed Successfully</h2>"
                f"<p><b>Actions Taken:</b> {len(tool_calls)}</p>"
                f"<h3>Executed Steps</h3>"
                f"<ol>{steps_html}</ol>"
            )
            try:
                self.sf.update_case(case_id, {
                    "AI_Resolution_Status__c": "Execute Resolution",
                    "AI_Recommended_Action__c": "Executed Successfully",
                    "AI_Resolution_Notes__c": rich_notes,
                })
            except Exception:
                pass

        # On successful execution, append case number to matching knowledge article
        if not failures and self.vector_store:
            try:
                case = self.sf.get_case_by_id(case_id)
                case_type = case.get("Type") or "General"
                existing = self._find_existing_article(case_type)
                if existing:
                    case_number = case.get("CaseNumber", "")
                    content = existing["content"]
                    if case_number and case_number not in content:
                        content = content.replace(
                            "## Case References\n",
                            f"## Case References\n{case_number}, ",
                        )
                        self._save_knowledge(existing["id"], content, case_type)
                        log.info("knowledge_case_appended", article=existing["id"], case=case_number)
            except Exception as e:
                log.warning("knowledge_append_failed", error=str(e))

        log.info("execution_complete", case_id=case_id, actions=len(tool_calls), failures=len(failures))
        return {
            "status": "executed" if not failures else "partial_failure",
            "case_id": case_id,
            "actions_taken": len(tool_calls),
            "failures": len(failures),
            "execution_log": execution_log,
        }

    def revise_resolution(
        self,
        case_id: str,
        agent_action: str,
        agent_comments: str,
        previous_resolution: dict,
        knowledge_text: str,
        company_knowledge: str,
    ) -> dict:
        log.info("revising_resolution", case_id=case_id, action=agent_action)

        case = self.sf.get_case_by_id(case_id)

        prompt = REVISE_PROMPT.format(
            case_id=case_id,
            case_number=case.get("CaseNumber", ""),
            subject=case.get("Subject", ""),
            description=case.get("Description", ""),
            case_type=case.get("Type", ""),
            previous_resolution=json.dumps(previous_resolution, indent=2),
            agent_action=agent_action,
            agent_comments=agent_comments,
            knowledge_articles=knowledge_text,
            company_knowledge=company_knowledge or "No company guidelines available.",
        )

        response = self.llm.messages.create(
            model=settings.anthropic_model,
            max_tokens=4096,
            messages=[{"role": "user", "content": prompt}],
        )

        import re

        raw = response.content[0].text.strip()
        match = re.search(r'\{[\s\S]*\}', raw)
        if not match:
            raise ValueError(f"No JSON found in AI response: {raw[:200]}")
        revised = json.loads(match.group())

        status_map = {
            "partially_correct": "Partially Correct Resolution",
            "not_right": "Incorrect Resolution",
        }

        steps_html = "".join(f"<li>{s}</li>" for s in revised.get("resolution_steps", []))
        rich_notes = (
            f"<h2>Revised AI Resolution</h2>"
            f"<p><b>Issue:</b> {revised.get('issue_summary', '')}</p>"
            f"<p><b>Root Cause:</b> {revised.get('root_cause', '')}</p>"
            f"<p><b>Confidence:</b> {revised.get('confidence_score', 0) * 100:.0f}%</p>"
            f"<h3>Resolution Steps</h3>"
            f"<ol>{steps_html}</ol>"
            f"<h3>Resolution Notes</h3>"
            f"<p>{revised.get('resolution_notes', '')}</p>"
            f"<h3>Customer Email</h3>"
            f"<p><b>Subject:</b> {revised.get('customer_email_subject', '')}</p>"
            f"<p>{revised.get('customer_email_body', '').replace(chr(10), '<br/>')}</p>"
        )

        try:
            self.sf.update_case(case_id, {
                "AI_Resolution_Notes__c": rich_notes,
                "AI_Confidence_Score__c": revised.get("confidence_score", 0),
                "AI_Recommended_Action__c": revised.get("recommended_action", ""),
                "AI_Resolution_Status__c": status_map.get(agent_action, "Pending"),
            })
        except Exception:
            log.warning("sf_update_fallback", case_id=case_id)

        self.tool_executor.execute_tool("add_case_comment", {
            "case_id": case_id,
            "comment": f"[AI Agent] Resolution revised.\nFeedback: {agent_comments}",
        })

        # Update knowledge base with agent's correction
        if self.vector_store:
            self._update_knowledge_from_feedback(
                case_id, case, agent_action, agent_comments, revised
            )

        log.info("resolution_revised", case_id=case_id)
        return revised

    def _update_knowledge_from_feedback(
        self, case_id: str, case: dict, agent_action: str,
        agent_comments: str, revised: dict,
    ):
        case_type = case.get("Type") or "General"
        case_number = case.get("CaseNumber", "")
        subject = case.get("Subject", "")

        # Find existing knowledge article for this case type
        existing_article = self._find_existing_article(case_type)

        if existing_article and agent_action == "not_right":
            # Incorrect — update existing article with corrected knowledge
            updated_content = self._llm_update_article(
                existing_content=existing_article["content"],
                existing_id=existing_article["id"],
                case_number=case_number,
                subject=subject,
                case_type=case_type,
                agent_comments=agent_comments,
                revised=revised,
                mode="replace_incorrect",
            )
            self._save_knowledge(existing_article["id"], updated_content, case_type)
        elif existing_article and agent_action == "partially_correct":
            # Partial — append correction to existing article
            updated_content = self._llm_update_article(
                existing_content=existing_article["content"],
                existing_id=existing_article["id"],
                case_number=case_number,
                subject=subject,
                case_type=case_type,
                agent_comments=agent_comments,
                revised=revised,
                mode="append_correction",
            )
            self._save_knowledge(existing_article["id"], updated_content, case_type)
        else:
            # No existing article — create new one
            new_content = self._llm_update_article(
                existing_content="",
                existing_id="",
                case_number=case_number,
                subject=subject,
                case_type=case_type,
                agent_comments=agent_comments,
                revised=revised,
                mode="create_new",
            )
            safe_type = re.sub(r'[^a-zA-Z0-9_-]', '_', case_type).strip('_')
            article_id = f"{safe_type}__agent_corrected"
            self._save_knowledge(article_id, new_content, case_type)

        log.info("knowledge_updated", case=case_number, action=agent_action)

    def _find_existing_article(self, case_type: str) -> dict | None:
        articles_dir = settings.knowledge_base_dir
        if not articles_dir.exists():
            return None

        safe_type = re.sub(r'[^a-zA-Z0-9_-]', '_', case_type).strip('_').lower()
        for md_file in articles_dir.glob("*.md"):
            if safe_type in md_file.stem.lower() and "correction_" not in md_file.stem:
                return {
                    "id": md_file.stem,
                    "content": md_file.read_text(encoding="utf-8"),
                }

        # Fallback: search vector store
        if self.vector_store:
            results = self.vector_store.search(case_type, top_k=1)
            if results and results[0]["similarity"] > 0.6:
                article_id = results[0]["id"]
                md_path = articles_dir / f"{article_id}.md"
                content = md_path.read_text(encoding="utf-8") if md_path.exists() else results[0]["content"]
                return {"id": article_id, "content": content}

        return None

    def _llm_update_article(
        self, existing_content: str, existing_id: str,
        case_number: str, subject: str, case_type: str,
        agent_comments: str, revised: dict, mode: str,
    ) -> str:
        steps = "\n".join(f"  {i+1}. {s}" for i, s in enumerate(revised.get("resolution_steps", [])))

        if mode == "create_new":
            instruction = (
                "Create a NEW knowledge article from the agent-corrected resolution below. "
                "Include the case number in Case References."
            )
        elif mode == "replace_incorrect":
            instruction = (
                "The existing knowledge article had INCORRECT resolution for this type of case. "
                "Update the article with the CORRECTED resolution from the agent. "
                "Add this case number to Case References. "
                "Replace incorrect resolution steps with the corrected ones. "
                "Keep any other valid content in the article."
            )
        else:
            instruction = (
                "The existing knowledge article is PARTIALLY correct. "
                "Append the agent's correction as additional resolution guidance. "
                "Add this case number to Case References. "
                "Do NOT remove existing content — merge the correction into the article."
            )

        prompt = f"""\
{instruction}

Case Number: {case_number}
Case Subject: {subject}
Case Type: {case_type}
Agent Feedback: {agent_comments}

Corrected Resolution Steps:
{steps}

Resolution Notes: {revised.get('resolution_notes', '')}

{"--- EXISTING ARTICLE ---" + chr(10) + existing_content + chr(10) + "--- END EXISTING ARTICLE ---" if existing_content else ""}

Generate the updated knowledge article in this markdown format:
# [Title]
## Case References
## Problem Statement
## Symptoms
## Root Cause
## Resolution Steps
## Customer Email Templates
## Metadata

Keep it concise. Resolution steps should be short actionable instructions.
Output ONLY the article markdown, no extra text."""

        response = self.llm.messages.create(
            model=settings.anthropic_model,
            max_tokens=4096,
            messages=[{"role": "user", "content": prompt}],
        )
        return response.content[0].text.strip()

    def _save_knowledge(self, article_id: str, content: str, case_type: str):
        articles_dir = settings.knowledge_base_dir
        articles_dir.mkdir(parents=True, exist_ok=True)

        # Save markdown
        md_path = articles_dir / f"{article_id}.md"
        md_path.write_text(content, encoding="utf-8")

        # Save docx
        self._save_docx(articles_dir / f"{article_id}.docx", content, case_type)

        # Update vector store
        if self.vector_store:
            self.vector_store.add_article(
                article_id=article_id,
                content=content,
                metadata={"case_type": case_type, "sub_type": "agent_corrected"},
            )

        # Update metadata
        meta_path = articles_dir / f"{article_id}.meta.json"
        existing_meta = {}
        if meta_path.exists():
            existing_meta = json.loads(meta_path.read_text(encoding="utf-8"))

        case_numbers = existing_meta.get("case_numbers", [])
        existing_meta.update({
            "case_type": case_type,
            "sub_type": "agent_corrected",
            "case_count": len(case_numbers),
            "case_numbers": case_numbers,
            "source": "agent_feedback",
        })
        meta_path.write_text(json.dumps(existing_meta, indent=2), encoding="utf-8")

        log.info("knowledge_saved", article_id=article_id)

    @staticmethod
    def _save_docx(path, content: str, case_type: str):
        doc = DocxDocument()

        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        for level, size, color in [
            ("Heading 1", 18, RGBColor(0x1B, 0x4F, 0x72)),
            ("Heading 2", 14, RGBColor(0x2E, 0x75, 0xB6)),
            ("Heading 3", 12, RGBColor(0x34, 0x98, 0xDB)),
        ]:
            s = doc.styles[level]
            s.font.name = "Arial"
            s.font.size = Pt(size)
            s.font.color.rgb = color

        header_para = doc.sections[0].header.paragraphs[0]
        run = header_para.add_run(f"Knowledge Article  |  {case_type}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        for line in content.split("\n"):
            stripped = line.strip()
            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif re.match(r"^\d+\.\s", stripped):
                text = re.sub(r"^\d+\.\s", "", stripped)
                doc.add_paragraph(text, style="List Number")
            elif stripped == "":
                continue
            else:
                doc.add_paragraph(stripped)

        doc.save(str(path))
