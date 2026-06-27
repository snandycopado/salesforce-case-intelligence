import json
import re
from collections import defaultdict
from pathlib import Path

import anthropic
import structlog
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import settings
from salesforce_client import SalesforceClient
from vector_store import VectorStore
from company_knowledge_loader import CompanyKnowledgeLoader

log = structlog.get_logger()

ARTICLE_GENERATION_PROMPT = """\
You are a Knowledge Base Author for a customer support organization.

Analyze the following group of resolved Salesforce Cases and produce a structured Knowledge Article.

Case Type: {case_type}
Sub Type: {sub_type}
Number of Cases: {case_count}

--- CASE DATA ---
{case_data}
--- END CASE DATA ---

Generate a Knowledge Article in the following format. Focus on actionable resolution knowledge.
Exclude unnecessary chat/communication details — only keep what helps resolve future cases. Resolution should be a 
list of action that AI agent will perform to resolve the case - Its could be
1. Send Email with following template. Follow the email template guide line document
2. Unlock user account - UPdate User account record by using MCP tools
3. Invoke Okta Sync - INvoke Okta INtegration to resync user data
4. Financial Transaction mismatch - Get All Financial Transaction detail and Ask LLM to find the root cause
5. Do not need follow up email
6. Resolution need should follow a logical sequence like end of resolution email will be send to customer
7. Resolution steps will feed to AI agent to resolve with help of other agents and MCP tool, So make the instruction very clear for agentic AI.
Always refer to the Company Standard Knowledge Guidelines below before drafting the final resolution.
The company guidelines define standard procedures, email contacts, templates, and escalation paths.
Your resolution steps MUST align with these company standards.

--- COMPANY STANDARD KNOWLEDGE ---
{company_knowledge}
--- END COMPANY STANDARD KNOWLEDGE ---

FORMAT:
# [Meaningful Title Describing the Problem]

## Case References
[List all CaseNumbers separated by commas]

## Problem Statement
[Clear description of the issue customers face]

## Symptoms
[Bullet list of observable symptoms]

## Root Cause
[What causes this issue]

## Resolution Steps
1. [Step-by-step resolution procedure]

## Troubleshooting Guide
[Additional diagnostic steps if initial resolution fails]


### Resolution Email
Subject: [subject line]
[email body]

### Follow-up Email
Subject: [subject line]
[email body]

## Metadata
- Average Resolution Time: [derived from data]
- Priority Distribution: [from data]
- Common Origin: [from data]
"""

MAX_ARTICLE_CASES = 30


class KnowledgeGenerator:
    def __init__(self, sf_client: SalesforceClient, vector_store: VectorStore):
        self.sf = sf_client
        self.vector_store = vector_store
        self.llm = anthropic.Anthropic(api_key=settings.anthropic_api_key)
        self.company_kb = CompanyKnowledgeLoader()
        settings.knowledge_base_dir.mkdir(parents=True, exist_ok=True)

    def generate_knowledge_base(self):
        log.info("starting_knowledge_generation")
        cases = self.sf.get_historical_cases()
        grouped = self._group_cases(cases)
        article_count = 0

        for (case_type, sub_type), case_list in grouped.items():
            chunks = [
                case_list[i : i + MAX_ARTICLE_CASES]
                for i in range(0, len(case_list), MAX_ARTICLE_CASES)
            ]

            for idx, chunk in enumerate(chunks):
                suffix = f"_part{idx + 1}" if len(chunks) > 1 else ""
                article = self._generate_article(case_type, sub_type, chunk)
                article_id = self._save_article(
                    case_type, sub_type, suffix, article, chunk
                )
                self._vectorize_article(article_id, article, case_type, sub_type)
                article_count += 1

        log.info("knowledge_generation_complete", articles=article_count)
        return article_count

    def _group_cases(self, cases: list[dict]) -> dict[tuple, list]:
        grouped = defaultdict(list)
        for case in cases:
            case_type = case.get("Type") or "General"
            sub_type = case.get("SubType__c") or case.get("Reason") or "General"
            grouped[(case_type, sub_type)].append(case)
        return grouped

    def _generate_article(
        self, case_type: str, sub_type: str, cases: list[dict]
    ) -> str:
        case_data = self._format_cases_for_prompt(cases)
        company_knowledge = self.company_kb.get_for_case_type(case_type)
        if not company_knowledge:
            company_knowledge = "No company standard guidelines available for this case type."

        response = self.llm.messages.create(
            model=settings.anthropic_model,
            max_tokens=4096,
            messages=[
                {
                    "role": "user",
                    "content": ARTICLE_GENERATION_PROMPT.format(
                        case_type=case_type,
                        sub_type=sub_type,
                        case_count=len(cases),
                        case_data=case_data,
                        company_knowledge=company_knowledge,
                    ),
                }
            ],
        )

        return response.content[0].text

    def _format_cases_for_prompt(self, cases: list[dict]) -> str:
        parts = []
        for case in cases:
            comments_text = "\n".join(
                f"  - [{c.get('CreatedDate', '')}] {c.get('CommentBody', '')}"
                for c in case.get("Comments", [])
            )
            emails_text = "\n".join(
                f"  - [{e.get('CreatedDate', '')}] From:{e.get('FromAddress', '')} "
                f"Subject:{e.get('Subject', '')} Body:{e.get('TextBody', '')[:300]}"
                for e in case.get("Emails", [])
            )

            parts.append(
                f"Case #{case.get('CaseNumber', 'N/A')}\n"
                f"  Subject: {case.get('Subject', '')}\n"
                f"  Description: {(case.get('Description') or '')[:500]}\n"
                f"  Status: {case.get('Status', '')}\n"
                f"  Type: {case.get('Type', '')}\n"
                f"  Priority: {case.get('Priority', '')}\n"
                f"  Origin: {case.get('Origin', '')}\n"
                f"  Created: {case.get('CreatedDate', '')}\n"
                f"  Closed: {case.get('ClosedDate', '')}\n"
                f"  Owner: {(case.get('Owner') or {}).get('Name', '')}\n"
                f"  Comments:\n{comments_text}\n"
                f"  Emails:\n{emails_text}\n"
            )
        return "\n---\n".join(parts)

    def _save_article(
        self,
        case_type: str,
        sub_type: str,
        suffix: str,
        content: str,
        cases: list[dict],
    ) -> str:
        safe_type = self._sanitize_filename(case_type)
        safe_sub = self._sanitize_filename(sub_type)
        article_id = f"{safe_type}__{safe_sub}{suffix}"

        # Save markdown
        filepath = settings.knowledge_base_dir / f"{article_id}.md"
        filepath.write_text(content, encoding="utf-8")

        # Save Word document
        self._save_as_docx(article_id, content, case_type, sub_type, cases)

        # Save metadata
        meta_path = settings.knowledge_base_dir / f"{article_id}.meta.json"
        meta = {
            "case_type": case_type,
            "sub_type": sub_type,
            "case_count": len(cases),
            "case_numbers": [c.get("CaseNumber") for c in cases],
        }
        meta_path.write_text(json.dumps(meta, indent=2), encoding="utf-8")

        log.info("article_saved", article_id=article_id, cases=len(cases))
        return article_id

    def _save_as_docx(
        self,
        article_id: str,
        content: str,
        case_type: str,
        sub_type: str,
        cases: list[dict],
    ):
        doc = DocxDocument()

        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        for level, size, color in [
            ("Heading 1", 18, RGBColor(0x1B, 0x4F, 0x72)),
            ("Heading 2", 14, RGBColor(0x2E, 0x75, 0xB6)),
            ("Heading 3", 12, RGBColor(0x34, 0x98, 0xDB)),
        ]:
            s = doc.styles[level]
            s.font.name = "Arial"
            s.font.size = Pt(size)
            s.font.color.rgb = color

        # Header with case type info
        header_para = doc.sections[0].header.paragraphs[0]
        header_run = header_para.add_run(f"Knowledge Article  |  {case_type} / {sub_type}")
        header_run.font.size = Pt(8)
        header_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Parse markdown content into Word
        for line in content.split("\n"):
            stripped = line.strip()

            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif re.match(r"^\d+\.\s", stripped):
                text = re.sub(r"^\d+\.\s", "", stripped)
                doc.add_paragraph(text, style="List Number")
            elif stripped.startswith("Subject:"):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(11)
            elif stripped == "":
                continue
            else:
                doc.add_paragraph(stripped)

        docx_path = settings.knowledge_base_dir / f"{article_id}.docx"
        doc.save(str(docx_path))
        log.info("docx_saved", path=str(docx_path))

    def _vectorize_article(
        self, article_id: str, content: str, case_type: str, sub_type: str
    ):
        self.vector_store.add_article(
            article_id=article_id,
            content=content,
            metadata={"case_type": case_type, "sub_type": sub_type},
        )

    @staticmethod
    def _sanitize_filename(name: str) -> str:
        return "".join(c if c.isalnum() or c in "-_" else "_" for c in name).strip("_")
